"""
CV Generation Service for PortfolioAI
"""
import os
import logging
import asyncio
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

from backend.utils.file_utils import get_temp_file
from backend.services.groq_client import GroqClient

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class CVGenerator:
    """Service for generating CVs in various formats."""
    
    def __init__(self, groq_client: Optional[GroqClient] = None):
        """Initialize the CV generator with an optional Groq client."""
        self.groq_client = groq_client or GroqClient()
    
    async def generate_cv(self, cv_data: Dict[str, Any], output_format: str = "docx") -> str:
        """
        Generate a CV file in the specified format.
        
        Args:
            cv_data: Dictionary containing CV data
            output_format: Output format ('docx', 'pdf', or 'md')
            
        Returns:
            str: Path to the generated CV file
            
        Raises:
            ValueError: If the output format is not supported
            asyncio.TimeoutError: If the CV generation times out
            Exception: For other errors during CV generation
        """
        logger.info(f"Starting CV generation in {output_format} format")
        
        if output_format not in ["docx", "md", "pdf"]:
            error_msg = f"Unsupported output format: {output_format}"
            logger.error(error_msg)
            raise ValueError(f"{error_msg}. Must be 'docx', 'md', or 'pdf'.")
        
        markdown_content = None
        output_file = None
        
        try:
            # Set a timeout for the entire CV generation process
            try:
                # Try to generate content with Groq first
                try:
                    logger.info("Attempting to generate CV content using Groq")
                    markdown_content = await asyncio.wait_for(
                        self.groq_client.generate_cv(cv_data),
                        timeout=90  # 1.5 minutes for Groq
                    )
                    logger.info("Successfully generated CV content using Groq")
                    
                except asyncio.TimeoutError:
                    logger.warning("Groq API request timed out, falling back to local generation")
                    markdown_content = self._generate_fallback_cv(cv_data)
                    
                except Exception as groq_error:
                    logger.warning(f"Error using Groq API: {str(groq_error)}, falling back to local generation")
                    markdown_content = self._generate_fallback_cv(cv_data)
                
                # Create the appropriate output file based on format
                output_file = get_temp_file(f".{output_format}")
                
                try:
                    logger.info(f"Creating CV in {output_format} format at {output_file}")
                    
                    if output_format == "docx":
                        self._create_docx(markdown_content, output_file)
                    elif output_format == "pdf":
                        self._create_pdf(markdown_content, output_file)
                    else:  # md
                        with open(output_file, 'w', encoding='utf-8') as f:
                            f.write(markdown_content)
                    
                    if not os.path.exists(output_file) or os.path.getsize(output_file) == 0:
                        raise RuntimeError(f"Failed to generate {output_format} file: Output file is empty or not created")
                    
                    logger.info(f"Successfully generated CV in {output_format} format at {output_file}")
                    return output_file
                    
                except Exception as format_error:
                    logger.error(f"Error generating {output_format} file: {str(format_error)}")
                    # Clean up the output file if it was created but corrupted
                    if output_file and os.path.exists(output_file):
                        try:
                            os.remove(output_file)
                            output_file = None
                        except Exception as cleanup_error:
                            logger.warning(f"Failed to clean up file {output_file}: {str(cleanup_error)}")
                    
                    # Try fallback to markdown if other formats fail
                    if output_format != "md":
                        logger.info("Attempting fallback to markdown format")
                        output_file = get_temp_file(".md")
                        with open(output_file, 'w', encoding='utf-8') as f:
                            f.write(markdown_content)
                        return output_file
                    
                    raise RuntimeError(f"Failed to generate CV: {str(format_error)}")
                
            except asyncio.TimeoutError as timeout_error:
                logger.error("CV generation timed out")
                raise RuntimeError("CV generation timed out. Please try again with a smaller CV or different format.")
                
        except Exception as e:
            logger.error(f"Unexpected error in CV generation: {str(e)}")
            raise RuntimeError(f"Failed to generate CV: {str(e)}")
            
        finally:
            # Clean up any temporary files if an error occurred
            if output_file and os.path.exists(output_file) and (not os.path.getsize(output_file) or not os.path.exists(output_file)):
                try:
                    os.remove(output_file)
                except Exception as cleanup_error:
                    logger.warning(f"Failed to clean up file {output_file}: {str(cleanup_error)}")
                    
    def _generate_fallback_cv(self, cv_data: Dict[str, Any]) -> str:
        """
        Generate a simple CV markdown as a fallback when Groq API is not available.
        
        Args:
            cv_data: Dictionary containing CV data
            
        Returns:
            str: Generated CV in markdown format
        """
        try:
            personal = cv_data.get("personal_info", {})
            work_exp = cv_data.get("work_experience", [])
            education = cv_data.get("education", [])
            skills = cv_data.get("skills", [])
            
            # Build the markdown content
            lines = []
            
            # Header
            lines.append(f"# {personal.get('name', 'Your Name')}")
            lines.append(f"{personal.get('email', '')} | {personal.get('phone', '')}")
            if personal.get('linkedin'):
                lines.append(f"LinkedIn: {personal['linkedin']}")
            if personal.get('github'):
                lines.append(f"GitHub: {personal['github']}")
            if personal.get('location'):
                lines.append(f"Location: {personal['location']}")
            
            lines.append("\n## Summary")
            lines.append(personal.get('summary', 'Experienced professional with a strong background in their field.'))
            
            # Work Experience
            if work_exp:
                lines.append("\n## Work Experience")
                for job in work_exp:
                    lines.append(f"\n### {job.get('title', 'Position')} at {job.get('company', 'Company')}")
                    dates = job.get('start_date', 'Start')
                    if job.get('end_date'):
                        dates += f" to {job['end_date']}"
                    elif job.get('current', False):
                        dates += " to Present"
                    lines.append(f"*{dates}*")
                    if job.get('location'):
                        lines.append(f"*{job['location']}*")
                    if job.get('description'):
                        lines.append("")
                        if isinstance(job['description'], list):
                            for desc in job['description']:
                                lines.append(f"- {desc}")
                        else:
                            lines.append(f"- {job['description']}")
            
            # Education
            if education:
                lines.append("\n## Education")
                for edu in education:
                    degree = edu.get('degree', 'Degree')
                    institution = edu.get('institution', 'Institution')
                    field = f" in {edu['field_of_study']}" if edu.get('field_of_study') else ""
                    lines.append(f"\n### {degree}{field}")
                    lines.append(f"{institution}")
                    dates = edu.get('start_date', 'Start')
                    if edu.get('end_date'):
                        dates += f" to {edu['end_date']}"
                    if 'gpa' in edu:
                        dates += f" | GPA: {edu['gpa']}"
                    lines.append(f"*{dates}*")
            
            # Skills
            if skills:
                lines.append("\n## Skills")
                if isinstance(skills, list):
                    lines.append(", ".join(skills))
                else:
                    lines.append(str(skills))
            
            return "\n".join(lines)
            
        except Exception as e:
            logger.error(f"Error in fallback CV generation: {str(e)}")
            return "# CV Generation Error\n\nUnable to generate CV. Please try again later or check your input data."
    
    def _add_heading(self, doc, text: str, level: int) -> None:
        """Add a heading to the document."""
        try:
            heading = doc.add_heading(level=level)
            run = heading.add_run(text)
            
            # Style the heading
            if level == 1:
                run.font.size = Pt(16)
                run.bold = True
            elif level == 2:
                run.font.size = Pt(14)
            elif level == 3:
                run.font.size = Pt(12)
                
        except Exception as e:
            logger.warning(f"Error adding heading to document: {str(e)}")
            raise
    
    def _create_docx(self, markdown_content: str, output_path: str) -> None:
        """Create a DOCX file from markdown content."""
        try:
            doc = Document()
            
            # Split the markdown into lines and process each one
            lines = markdown_content.split('\n')
            current_paragraph = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    current_paragraph = None
                    continue
                        
                # Handle headings
                if line.startswith('### '):
                    self._add_heading(doc, line[4:], 3)
                    current_paragraph = None
                elif line.startswith('## '):
                    self._add_heading(doc, line[3:], 2)
                    current_paragraph = None
                elif line.startswith('# '):
                    self._add_heading(doc, line[2:], 1)
                    current_paragraph = None
                # Handle bullet points
                elif line.startswith('- '):
                    if current_paragraph is None:
                        current_paragraph = doc.add_paragraph(style='List Bullet')
                    current_paragraph.add_run(line[2:])
                # Regular paragraph
                else:
                    if current_paragraph is None:
                        current_paragraph = doc.add_paragraph()
                        current_paragraph.add_run(line)
                    else:
                        current_paragraph.add_run(' ' + line)
            
            doc.save(output_path)
            
        except Exception as e:
            logger.error(f"Error creating DOCX: {str(e)}")
            raise
            
    def _create_pdf(self, markdown_content: str, output_path: str) -> None:
        """Create a PDF file from markdown content."""
        try:
            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=72, 
                leftMargin=72,
                topMargin=72, 
                bottomMargin=72
            )
            
            styles = getSampleStyleSheet()
            elements = []
            
            # Add a custom style for normal text with better spacing
            normal_style = ParagraphStyle(
                name='NormalSpaced',
                parent=styles['Normal'],
                spaceAfter=6,
                fontSize=11
            )
            
            # Split the markdown into lines and process each one
            lines = markdown_content.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    elements.append(Spacer(1, 12))  # Add some space between paragraphs
                    continue
                        
                # Handle headings
                if line.startswith('### '):
                    elements.append(Paragraph(f"<b>{line[4:]}</b>", styles['Heading3']))
                elif line.startswith('## '):
                    elements.append(Paragraph(f"<u>{line[3:]}</u>", styles['Heading2']))
                elif line.startswith('# '):
                    elements.append(Paragraph(line[2:], styles['Title']))
                # Handle bullet points
                elif line.startswith('- '):
                    elements.append(Paragraph(f"• {line[2:]}", normal_style))
                # Regular paragraph
                else:
                    elements.append(Paragraph(line, normal_style))
                
                elements.append(Spacer(1, 6))  # Add small space after each line
            
            doc.build(elements)
            
        except Exception as e:
            logger.error(f"Error creating PDF: {str(e)}")
            raise

# Singleton instance
cv_generator = CVGenerator()
