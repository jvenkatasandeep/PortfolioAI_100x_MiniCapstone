import streamlit as st
import os
import io
import asyncio
import logging
import requests
from dotenv import load_dotenv
from typing import Dict, Optional, List, Any, Tuple
from pathlib import Path
from PyPDF2 import PdfReader
import docx
import pyperclip
from components.header import show_header

# Load environment variables from .env if available
load_dotenv()

# Backend API URL
# Allow override via environment variable or Streamlit secrets
BACKEND_URL = os.environ.get(
    "BACKEND_URL",
    st.secrets.get("BACKEND_URL", "http://localhost:8000") if hasattr(st, "secrets") else "http://localhost:8000",
)

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Session state keys
PORTFOLIO_DATA = 'portfolio_data'

# Available portfolio sections
DEFAULT_SECTIONS = {
    'about': 'About Me',
    'experience': 'Work Experience',
    'education': 'Education',
    'skills': 'Skills',
    'projects': 'Projects',
    'contact': 'Contact'
}

# Default template for portfolio generation
PORTFOLIO_TEMPLATE = """
# {name}'s Portfolio

## About Me
{about}

## Work Experience
{experience}

## Education
{education}

## Skills
{skills}

## Projects
{projects}

## Contact
{contact}
"""

def initialize_session_state():
    """Initialize session state for portfolio generator."""
    if PORTFOLIO_DATA not in st.session_state:
        st.session_state[PORTFOLIO_DATA] = {
            'resume_text': '',
            'sections': {k: '' for k in DEFAULT_SECTIONS.keys()},
            'generated_portfolio': None
        }

def show_portfolio_header():
    """Display the portfolio generator specific header."""
    st.markdown("<h2 style='text-align: center; margin-bottom: 1rem;'>Portfolio Generator</h2>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: #666; margin-bottom: 1.5rem;'>Create a professional portfolio website based on your resume.</p>", unsafe_allow_html=True)

def process_resume(file) -> Optional[str]:
    """Process an uploaded resume file and extract text."""
    if file is None:
        return None
    
    try:
        file_ext = Path(file.name).suffix.lower()
        file_content = file.getvalue()
        
        if file_ext == '.pdf':
            pdf_reader = PdfReader(io.BytesIO(file_content))
            text = "\n\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
        elif file_ext in ['.docx', '.doc']:
            doc = docx.Document(io.BytesIO(file_content))
            text = "\n\n".join([para.text for para in doc.paragraphs if para.text])
        else:  # txt file
            text = file_content.decode("utf-8")
            
        return text.strip()
    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
        logger.error(f"File processing error: {str(e)}")
        return None

def show_resume_upload() -> bool:
    """Show resume upload section and return True if resume was processed."""
    st.subheader("1. Upload Your Resume")
    uploaded_file = st.file_uploader(
        "Upload your resume (PDF, DOCX, or TXT)",
        type=["pdf", "docx", "doc", "txt"]
    )
    
    if uploaded_file is not None:
        with st.spinner("Processing your resume..."):
            resume_text = process_resume(uploaded_file)
            if resume_text:
                st.session_state[PORTFOLIO_DATA]['resume_text'] = resume_text
                st.success("Resume processed successfully!")
                return True
    
    return bool(st.session_state[PORTFOLIO_DATA].get('resume_text'))

def show_section_selection() -> List[str]:
    """Show section selection checkboxes and return selected sections."""
    st.subheader("2. Select Portfolio Sections")
    st.markdown("Choose which sections to include in your portfolio:")
    
    selected_sections = []
    cols = st.columns(2)
    
    for i, (section_key, section_name) in enumerate(DEFAULT_SECTIONS.items()):
        with cols[i % 2]:
            # Add custom CSS to make checkbox text black
            st.markdown(
                f"""
                <style>
                    div[data-testid="stCheckbox"] label p {{
                        color: black !important;
                    }}
                </style>
                """,
                unsafe_allow_html=True
            )
            if st.checkbox(
                section_name,
                value=True,
                key=f"section_{section_key}",
                help=f"Include {section_name.lower()} section"
            ):
                selected_sections.append(section_key)
    
    if not selected_sections:
        st.warning("Please select at least one section.")
    
    return selected_sections

def show_generated_portfolio(portfolio_data: Optional[Dict]):
    """Display the generated portfolio.
    
    Args:
        portfolio_data: Dictionary containing the portfolio data or None
    """
    if not portfolio_data:
        st.warning("No portfolio data available to display.")
        return
        
    st.subheader("Your Portfolio")
    
    # Show public URL if subdomain is available
    if 'subdomain' in portfolio_data and portfolio_data['subdomain']:
        public_url = f"http://{portfolio_data['subdomain']}.portfolioai.com"
        st.success("🎉 Your portfolio has been published!")
        st.markdown(f"### 🌐 Public URL: [{public_url}]({public_url})")
        
        # Add a copy to clipboard button
        if st.button("📋 Copy URL to Clipboard"):
            import pyperclip
            pyperclip.copy(public_url)
            st.toast("URL copied to clipboard!")
        
        st.markdown("---")
    
    try:
        if not isinstance(portfolio_data, dict):
            st.error("Invalid portfolio data format.")
            logger.error(f"Invalid portfolio data type: {type(portfolio_data)}")
            return
            
        if 'content' in portfolio_data and portfolio_data['content']:
            # If we have HTML content, render it
            st.components.v1.html(portfolio_data['content'], height=800, scrolling=True)
        else:
            # Fallback to showing raw data if content is not available
            st.warning("No content found in the portfolio data. Showing raw data instead.")
            for section, content in portfolio_data.items():
                if section in ['sections', 'subdomain', 'public_url']:
                    continue  # Skip internal keys
                    
                with st.expander(section.replace('_', ' ').title()):
                    if isinstance(content, str):
                        st.markdown(content)
                    elif isinstance(content, (list, dict)):
                        st.json(content)
                    else:
                        st.text(str(content))
    except Exception as e:
        st.error(f"Error displaying portfolio: {str(e)}")
        logger.exception("Error in show_generated_portfolio")

async def call_portfolio_api(resume_text: str, sections: List[str], subdomain: str = '') -> Tuple[bool, Dict]:
    """
    Call the backend API to generate portfolio content with optional subdomain.
    
    Args:
        resume_text: The text content of the resume
        sections: List of sections to include in the portfolio
        subdomain: Optional subdomain for the portfolio URL
        
    Returns:
        Tuple of (success: bool, response_data: dict)
    """
    try:
        # Set up headers with API key if available
        headers = {
            "Content-Type": "application/json"
        }
        api_key = os.getenv("PORTFOLIOAI_API_KEY")
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"
        
        # Prepare request data with optional subdomain
        request_data = {
            "resume_text": resume_text,
            "sections": sections
        }
        
        # Add subdomain to request if provided
        if subdomain:
            request_data["subdomain"] = subdomain.lower()  # Ensure lowercase
        
        response = requests.post(
            f"{BACKEND_URL}/api/portfolio/generate/from-resume",
            headers=headers,
            json=request_data,
            timeout=300  # 5 minutes timeout
        )
        
        if response.status_code == 200:
            return True, response.json()
        else:
            error_msg = f"API Error: {response.status_code} - {response.text}"
            logger.error(error_msg)
            return False, {"error": error_msg}
            
    except requests.exceptions.RequestException as e:
        error_msg = f"Request failed: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return False, {"error": error_msg}

async def generate_portfolio_content(resume_text: str, sections: List[str], subdomain: str = '') -> Optional[Dict]:
    """Generate portfolio content using the backend API with optional subdomain."""
    with st.spinner("Generating portfolio content with AI..."):
        success, result = await call_portfolio_api(resume_text, sections, subdomain)
        
        if success and result.get('status') == 'success':
            st.success("Portfolio generated successfully!")
            # Add subdomain to the result if available
            return {
                'status': 'success',
                'content': result.get('content', ''),
                'sections': sections,
                'model': result.get('model', 'unknown'),
                'usage': result.get('usage', {}),
                'subdomain': result.get('subdomain', '')
            }
        else:
            error_msg = result.get('error', 'Failed to generate portfolio')
            st.error(f"Error: {error_msg}")
            logger.error(f"Portfolio generation failed: {error_msg}")
            return None

def check_backend_connection() -> bool:
    """Check if the backend API is reachable."""
    try:
        response = requests.get(f"{BACKEND_URL}/health")
        return response.status_code == 200
    except requests.exceptions.RequestException:
        return False

def check_subdomain_available(subdomain: str) -> Tuple[bool, str]:
    """Check if a subdomain is available.
    
    Returns:
        Tuple of (is_available, message)
    """
    try:
        response = requests.get(
            f"{BACKEND_URL}/api/portfolio/check-subdomain",
            params={"subdomain": subdomain}
        )
        
        if response.status_code == 200:
            data = response.json()
            return data.get('available', False), data.get('message', '')
            
        return False, f"Error checking subdomain: {response.status_code}"
        
    except requests.exceptions.RequestException as e:
        return False, f"Failed to check subdomain: {str(e)}"

def show_portfolio_generator():
    """Main function to display the portfolio generator."""
    # Initialize session state
    if 'resume_processed' not in st.session_state:
        st.session_state['resume_processed'] = False
    
    if PORTFOLIO_DATA not in st.session_state:
        st.session_state[PORTFOLIO_DATA] = {}
    
    # Show portfolio header
    st.title("Portfolio Generator")
    
    # Only show uploader if resume not yet processed
    if not st.session_state.get('resume_processed', False):
        st.write("### 1. Upload Your Resume")
        uploaded_file = st.file_uploader("Choose a resume file", type=["pdf", "docx", "txt"])
        
        if uploaded_file is not None:
            try:
                # Process the file
                if uploaded_file.name.endswith('.pdf'):
                    pdf_reader = PdfReader(uploaded_file)
                    text = "\n\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
                elif uploaded_file.name.endswith(('.docx', '.doc')):
                    doc = docx.Document(uploaded_file)
                    text = "\n\n".join([para.text for para in doc.paragraphs if para.text])
                else:  # txt file
                    text = uploaded_file.getvalue().decode("utf-8")
                
                if text.strip():
                    st.session_state[PORTFOLIO_DATA]['resume_text'] = text.strip()
                    st.session_state['resume_processed'] = True
                    st.session_state['current_step'] = 'subdomain'  # Move to next step
                    st.rerun()
                else:
                    st.error("The uploaded file appears to be empty.")
                    
            except Exception as e:
                st.error(f"Error processing file: {str(e)}")
                logger.error(f"File processing error: {str(e)}")
        
        # Stop here if we haven't processed a resume yet
        return
    
    # Subdomain validation function
    def is_valid_subdomain(subdomain: str) -> bool:
        import re
        pattern = r'^[a-z0-9]([a-z0-9\-]{0,61}[a-z0-9])?$'
        return bool(re.match(pattern, subdomain)) and 3 <= len(subdomain) <= 63
    
    # If we get here, we've processed the resume
    if st.session_state.get('current_step') == 'portfolio' and PORTFOLIO_DATA in st.session_state and st.session_state[PORTFOLIO_DATA].get('generated_portfolio'):
        # Show the generated portfolio
        st.subheader("Your Generated Portfolio")
        portfolio_data = st.session_state[PORTFOLIO_DATA]['generated_portfolio']
        if isinstance(portfolio_data, dict) and portfolio_data.get('status') == 'success':
            show_generated_portfolio(portfolio_data)
        else:
            st.error("Failed to generate portfolio. Please try again.")
            logger.error(f"Invalid portfolio data: {portfolio_data}")
        return
    elif st.session_state.get('current_step') == 'subdomain':
        st.success("✓ Resume uploaded successfully!")
        st.write("### 2. Choose Your Portfolio URL")
        st.markdown("Select a custom subdomain for your portfolio (e.g., 'john-doe' for 'john-doe.portfolioai.com')")
    
        col1, col2 = st.columns([3, 1])
        with col1:
            subdomain = st.text_input(
                "Choose a subdomain", 
                value=st.session_state[PORTFOLIO_DATA].get('subdomain', ''),
                max_chars=63,
                help="3-63 characters, lowercase letters, numbers, and hyphens only. Must start and end with a letter or number.",
                key="subdomain_input"
            ).lower()  # Convert to lowercase automatically
        
        with col2:
            st.markdown("\n\n")  # For alignment
            st.markdown(".portfolioai.com")
        
        # Validate subdomain format
        if subdomain:
            if not is_valid_subdomain(subdomain):
                st.error("❌ Invalid subdomain format. Use 3-63 characters: lowercase letters, numbers, and hyphens only. Must start and end with a letter or number.")
                st.session_state['subdomain_available'] = False
                st.stop()
            
            # Check subdomain availability
            with st.spinner("Checking subdomain availability..."):
                is_available, message = check_subdomain_available(subdomain)
                
                if is_available:
                    st.success(f"✅ Subdomain available! Your portfolio will be at: http://{subdomain}.portfolioai.com")
                    st.session_state['subdomain_available'] = True
                else:
                    st.error(f"❌ {message} Please choose a different subdomain.")
                    st.session_state['subdomain_available'] = False
                    st.stop()
            
            # Store valid subdomain in session
            st.session_state[PORTFOLIO_DATA]['subdomain'] = subdomain
        else:
            st.warning("ℹ️ Please choose a subdomain for your portfolio")
            st.session_state['subdomain_available'] = False
            st.stop()  # Stop execution if no subdomain is provided
        
        # Get resume text from session state
        resume_text = st.session_state[PORTFOLIO_DATA].get('resume_text', '')
        if not resume_text:
            st.warning("Please upload a valid resume first.")
            st.session_state['resume_processed'] = False
            return
        
        # Show section selection
        st.subheader("3. Select Sections to Include")
        selected_sections = show_section_selection()
        
        # Show generate button if sections are selected and subdomain is available
        if selected_sections and st.button("Generate Portfolio", disabled=not st.session_state.get('subdomain_available', False)):
            # Generate portfolio content
            with st.spinner("Generating your portfolio..."):
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                try:
                    # Get subdomain from session state
                    subdomain = st.session_state[PORTFOLIO_DATA].get('subdomain', '')
                    
                    # Call API with subdomain
                    portfolio_data = loop.run_until_complete(
                        generate_portfolio_content(resume_text, selected_sections, subdomain)
                    )
                    if portfolio_data:
                        st.session_state[PORTFOLIO_DATA]['generated_portfolio'] = portfolio_data
                        st.session_state['current_step'] = 'portfolio'  # Move to portfolio view
                        st.rerun()  # Rerun to show the generated portfolio
                except Exception as e:
                    st.error(f"An error occurred while generating the portfolio: {str(e)}")
                    logger.exception("Portfolio generation failed")
                finally:
                    loop.close()
        
        # After generating portfolio, we'll rerun to show the preview
        # The preview will be handled at the beginning of the function

# Run the portfolio generator
if __name__ == "__main__":
    show_portfolio_generator()